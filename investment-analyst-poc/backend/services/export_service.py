from __future__ import annotations
import io
from typing import Dict, Any
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document as DocxDocument


class ExportService:
    """Handles PDF and Excel export of analysis results."""

    def _citation_lines(self, value: Dict[str, Any]) -> list[str]:
        citations = value.get("citations") or []
        lines: list[str] = []
        for citation in citations[:3]:
            source_name = citation.get("source_name", "Source")
            source_uri = citation.get("source_uri", "")
            excerpt = citation.get("excerpt", "")
            lines.append(f"Source: {source_name}")
            if source_uri:
                lines.append(source_uri)
            if excerpt:
                lines.append(excerpt)
        return lines

    def export_pdf(self, results: Dict[str, Any]) -> bytes:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph("Investment Analysis Report", styles['Title']))
        story.append(Spacer(1, 12))
        for k, v in results.items():
            story.append(Paragraph(k.replace("_", " ").title(), styles['Heading2']))
            if isinstance(v, dict) and 'text' in v:
                story.append(Paragraph(str(v.get('text', '')), styles['BodyText']))
                if 'score' in v:
                    story.append(Paragraph(f"Score: {v['score']}/10", styles['BodyText']))
                if 'confidence' in v:
                    story.append(Paragraph(f"Confidence: {float(v['confidence']) * 100:.0f}%", styles['BodyText']))
                for line in self._citation_lines(v):
                    story.append(Paragraph(line, styles['BodyText']))
            else:
                story.append(Paragraph(str(v), styles['BodyText']))
            story.append(Spacer(1, 8))

        doc.build(story)
        return buf.getvalue()

    def export_docx(self, results: Dict[str, Any]) -> bytes:
        doc = DocxDocument()
        doc.add_heading('Investment Analysis Report', 0)
        for k, v in results.items():
            doc.add_heading(k.replace('_', ' ').title(), level=2)
            if isinstance(v, dict) and 'text' in v:
                doc.add_paragraph(str(v.get('text', '')))
                if 'score' in v:
                    doc.add_paragraph(f"Score: {v['score']}/10")
                if 'confidence' in v:
                    doc.add_paragraph(f"Confidence: {float(v['confidence']) * 100:.0f}%")
                for line in self._citation_lines(v):
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph(str(v))
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
