from __future__ import annotations
import os
import asyncio
from io import BytesIO
from typing import List
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from ..utils.logger import get_logger
from ..utils.helpers import sanitize_text

logger = get_logger(__name__)


class DocumentParser:
    """Parses PDF and DOCX documents asynchronously."""

    async def parse_bytes(self, file_name: str, data: bytes) -> str:
        ext = os.path.splitext(file_name)[1].lower()
        if ext == ".pdf":
            return await self._parse_pdf_bytes(data)
        if ext == ".docx":
            return await self._parse_docx_bytes(data)
        raise ValueError("Unsupported file type")

    async def parse_file(self, path: str) -> str:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            return await self._parse_pdf(path)
        if ext == ".docx":
            return await self._parse_docx(path)
        raise ValueError("Unsupported file type")

    async def _parse_pdf(self, path: str) -> str:
        def _sync() -> str:
            doc = fitz.open(path)
            parts: List[str] = []
            try:
                for page in doc:
                    parts.append(page.get_text("text"))
            finally:
                doc.close()
            return sanitize_text("\n".join(parts))
        return await asyncio.to_thread(_sync)

    async def _parse_pdf_bytes(self, data: bytes) -> str:
        def _sync() -> str:
            doc = fitz.open(stream=data, filetype="pdf")
            parts: List[str] = []
            try:
                for page in doc:
                    parts.append(page.get_text("text"))
            finally:
                doc.close()
            return sanitize_text("\n".join(parts))
        return await asyncio.to_thread(_sync)

    async def _parse_docx(self, path: str) -> str:
        def _sync() -> str:
            d = DocxDocument(path)
            parts = [p.text for p in d.paragraphs]
            return sanitize_text("\n".join(parts))
        return await asyncio.to_thread(_sync)

    async def _parse_docx_bytes(self, data: bytes) -> str:
        def _sync() -> str:
            d = DocxDocument(BytesIO(data))
            parts = [p.text for p in d.paragraphs]
            return sanitize_text("\n".join(parts))
        return await asyncio.to_thread(_sync)
